#!/usr/bin/env python3
"""Export resume/summary YAML data to DOCX format.

Supports two visual styles:
  - ats:    Clean, single-column, no colors/images. Optimized for ATS parsing.
  - styled: Accent-colored headings, photo, company logos, recommendation blocks.

Auto-detects whether the input is a *full resume* (has work.entries[].company)
or a *summary* (has clients/tech_stack) and renders accordingly.
"""

import argparse
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor


# ──────────────────────────────────────────────
# CONSTANTS
# ──────────────────────────────────────────────
ACCENT = RGBColor(0x1A, 0x52, 0x76)
MID_GRAY = RGBColor(0x5D, 0x6D, 0x7E)
LIGHT_GRAY = RGBColor(0x95, 0xA5, 0xA6)
DARK = RGBColor(0x1C, 0x28, 0x33)

ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"

MONTH_NAMES_DE = {
    1: "Jan.", 2: "Feb.", 3: "März", 4: "Apr.", 5: "Mai", 6: "Juni",
    7: "Juli", 8: "Aug.", 9: "Sep.", 10: "Okt.", 11: "Nov.", 12: "Dez.",
}
MONTH_NAMES_EN = {
    1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
    7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec",
}


# ──────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────
def clean_text(text: str) -> str:
    """Clean up YAML text: normalize whitespace, convert --- to em-dash."""
    if not text:
        return ""
    text = str(text).strip()
    text = text.replace("---", "\u2014")
    text = text.replace("--", "\u2013")
    # Collapse internal newlines from YAML block scalars into spaces
    text = re.sub(r"\s*\n\s*", " ", text)
    return text


def fmt_date(value: str, lang: str = "en") -> str:
    """Format YYYY-MM-DD or YYYY-MM to 'Mon YYYY' or just 'YYYY'."""
    if not value:
        return ""
    s = str(value).strip()
    months = MONTH_NAMES_DE if lang == "de" else MONTH_NAMES_EN
    for fmt in ("%Y-%m-%d", "%Y-%m"):
        try:
            dt = datetime.strptime(s, fmt)
            return f"{months[dt.month]} {dt.year}"
        except ValueError:
            continue
    if re.fullmatch(r"\d{4}", s):
        return s
    return s


def fmt_date_range(start: str, end: str, lang: str = "en") -> str:
    """Format a date range like 'May 2019 – Jul 2024' or 'May 2019 – present'."""
    s = fmt_date(start, lang)
    if end:
        e = fmt_date(end, lang)
    else:
        e = "heute" if lang == "de" else "present"
    if s and e:
        return f"{s} \u2013 {e}"
    return s or e


def fmt_year_range(start: str, end: str, lang: str = "en") -> str:
    """Format a year range like '2019 – 2024'."""
    s = str(start)[:4] if start else ""
    if end:
        e = str(end)[:4]
    else:
        e = "heute" if lang == "de" else "present"
    if s and e:
        return f"{s} \u2013 {e}"
    return s or e


def is_resume_data(data: dict) -> bool:
    """Detect if YAML data is a full resume (vs. summary)."""
    work = data.get("work", {})
    entries = work.get("entries", [])
    if entries and "company" in entries[0]:
        return True
    return False


def resolve_asset(filename: str) -> Path | None:
    """Resolve an asset filename to a full path, return None if missing."""
    if not filename:
        return None
    p = ASSETS_DIR / filename
    if p.exists():
        return p
    return None


def resolve_logo(filename: str) -> Path | None:
    """Resolve a logo filename, converting SVG/GIF to PDF if available."""
    if not filename:
        return None
    p = ASSETS_DIR / filename
    # python-docx can't embed SVG/GIF; try PNG/JPG fallback or PDF
    if p.suffix.lower() in (".svg", ".gif"):
        for alt_ext in (".png", ".jpg", ".jpeg"):
            alt = p.with_suffix(alt_ext)
            if alt.exists():
                return alt
        # Check if render_svg.py created a PDF version
        pdf = p.with_suffix(".pdf")
        if pdf.exists():
            return None  # python-docx can't use PDF images either
        return None
    if p.exists():
        return p
    return None


def set_cell_shading(cell, hex_color: str) -> None:
    """Set background fill color on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, **kwargs) -> None:
    """Set borders on a table cell. kwargs: top, bottom, start, end with value dicts."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, props in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for attr, val in props.items():
            el.set(qn(f"w:{attr}"), str(val))
        borders.append(el)
    tc_pr.append(borders)


def remove_table_borders(table) -> None:
    """Remove all borders from a table."""
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_vertical_alignment(cell, align: str = "center") -> None:
    """Set vertical alignment on a table cell ('top', 'center', 'bottom')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), align)
    tc_pr.append(v_align)


def set_cell_width(cell, width) -> None:
    """Set fixed width on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def add_accent_rule(doc: Document) -> None:
    """Add a horizontal accent-colored rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._element.get_or_add_pPr()
    bottom = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "12")
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), "1A5276")
    bottom.append(b)
    p_pr.append(bottom)


# ──────────────────────────────────────────────
# STYLE SETUP
# ──────────────────────────────────────────────
def setup_styles(doc: Document, styled: bool) -> None:
    """Configure document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK if styled else RGBColor(0x33, 0x33, 0x33)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = Pt(13)

    # Heading 1 — section headings
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT if styled else RGBColor(0x33, 0x33, 0x33)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — sub-entries (job title, product name)
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = DARK
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(1)
    h2.paragraph_format.keep_with_next = True

    # Create a "Quote" style for recommendations
    if "Quote" not in [s.name for s in doc.styles]:
        quote_style = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.name = "Calibri"
        quote_style.font.size = Pt(9)
        quote_style.font.italic = True
        quote_style.font.color.rgb = MID_GRAY if styled else RGBColor(0x55, 0x55, 0x55)
        quote_style.paragraph_format.left_indent = Cm(0.5)
        quote_style.paragraph_format.space_before = Pt(3)
        quote_style.paragraph_format.space_after = Pt(3)


# ──────────────────────────────────────────────
# SECTION RENDERERS — FULL RESUME
# ──────────────────────────────────────────────
def add_header_resume(doc: Document, data: dict, styled: bool) -> None:
    """Render the header block with name, title, and contact info."""
    basics = data["basics"]
    lang = data.get("meta", {}).get("language", "en")

    address = basics.get("address", {})
    addr_str = address.get("full", "") if isinstance(address, dict) else str(address)
    company = basics.get("company", "")
    title = basics.get("title", "")
    label = f"{company}\u00AE | {title}" if company else title

    if styled:
        # ── Accent banner table: [photo | name + title] ──
        photo_path = resolve_asset(basics.get("photo", ""))

        table = doc.add_table(rows=1, cols=2 if photo_path else 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)
        # Make table span full page width (edge to edge) by using negative indent
        tbl_pr = table._tbl.tblPr
        # Negative indent to pull table to left margin edge
        tbl_indent = OxmlElement("w:tblInd")
        tbl_indent.set(qn("w:w"), "-1134")  # -2 cm in twips
        tbl_indent.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_indent)
        # Set table width to page width (21cm - 0 margins → use larger than content area)
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")  # 100% of page
        # Tighten bottom cell margin only; keep top for image breathing room
        cell_mar = OxmlElement("w:tblCellMar")
        el = OxmlElement("w:bottom")
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
        tbl_pr.append(cell_mar)

        if photo_path:
            # Left cell: photo on accent background
            photo_cell = table.cell(0, 0)
            set_cell_shading(photo_cell, "1A5276")
            set_cell_vertical_alignment(photo_cell, "center")
            set_cell_width(photo_cell, 1700)
            # Bottom border as the darker accent strip
            set_cell_borders(photo_cell, bottom={
                "val": "single", "sz": "18", "space": "0", "color": "143D59",
            })
            pp = photo_cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(6)
            pp.paragraph_format.space_after = Pt(6)
            pp.paragraph_format.left_indent = Cm(0.3)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = pp.add_run()
            run.add_picture(str(photo_path), width=Cm(2.5))

            text_cell = table.cell(0, 1)
        else:
            text_cell = table.cell(0, 0)

        # Text cell: name + title on accent background
        set_cell_shading(text_cell, "1A5276")
        set_cell_vertical_alignment(text_cell, "center")
        # Bottom border as the darker accent strip
        set_cell_borders(text_cell, bottom={
            "val": "single", "sz": "18", "space": "0", "color": "143D59",
        })

        # Name
        np = text_cell.paragraphs[0]
        np.paragraph_format.space_before = Pt(10)
        np.paragraph_format.space_after = Pt(1)
        np.paragraph_format.left_indent = Cm(0.4)
        run = np.add_run(basics["name"])
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Company | Title
        tp = text_cell.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(10)
        tp.paragraph_format.left_indent = Cm(0.4)
        run = tp.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xD6, 0xE4, 0xEF)

        # Contact info row below banner — tight spacing
        contact_parts = [
            addr_str,
            basics.get("email", ""),
            basics.get("phone", ""),
            basics.get("website", {}).get("text", ""),
        ]
        contact_line = " | ".join(part for part in contact_parts if part)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(contact_line)
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

        # Profiles row
        profiles = basics.get("profiles", [])
        if profiles:
            profile_texts = []
            for pr in profiles:
                network = pr.get("network", "")
                text = pr.get("text", pr.get("url", ""))
                profile_texts.append(f"{network}: {text}" if network else text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(" | ".join(profile_texts))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY

    else:
        # ── ATS header: plain centered text ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(basics["name"])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DARK

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = MID_GRAY

        contact_parts = [
            addr_str,
            basics.get("email", ""),
            basics.get("phone", ""),
            basics.get("website", {}).get("text", ""),
        ]
        contact_line = " | ".join(part for part in contact_parts if part)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_line)
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

        profiles = basics.get("profiles", [])
        if profiles:
            profile_texts = []
            for pr in profiles:
                network = pr.get("network", "")
                text = pr.get("text", pr.get("url", ""))
                profile_texts.append(f"{network}: {text}" if network else text)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(profile_texts))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY


def add_summary_section(doc: Document, data: dict, styled: bool) -> None:
    """Render professional summary."""
    basics = data.get("basics", {})
    summary_text = basics.get("summary", "")
    if not summary_text:
        return
    lang = data.get("meta", {}).get("language", "en")
    title = "Profil" if lang == "de" else "Professional Summary"
    doc.add_heading(title, level=1)
    doc.add_paragraph(clean_text(summary_text))


def add_services_section(doc: Document, data: dict, styled: bool) -> None:
    """Render core services."""
    services = data.get("services")
    if not services:
        return
    doc.add_heading(services.get("title", "Services"), level=1)
    for entry in services.get("entries", []):
        p = doc.add_paragraph()
        run = p.add_run(entry.get("title", ""))
        run.bold = True
        desc = entry.get("description", "")
        if desc:
            p.add_run(f" \u2014 {clean_text(desc)}")


def add_skills_section(doc: Document, data: dict, styled: bool) -> None:
    """Render technical skills."""
    skills = data.get("skills")
    if not skills:
        return
    doc.add_heading(skills.get("title", "Skills"), level=1)

    for cat in skills.get("categories", []):
        label = cat.get("label", "")
        level = cat.get("level", "")
        keywords = cat.get("keywords", [])
        p = doc.add_paragraph()
        run = p.add_run(f"{label}")
        run.bold = True
        if level:
            run = p.add_run(f" ({level})")
            run.font.color.rgb = MID_GRAY
            run.font.size = Pt(9)
        p.add_run(f": {', '.join(keywords)}")

    # Human languages
    languages = data.get("languages")
    if languages:
        p = doc.add_paragraph()
        run = p.add_run(languages.get("title", "Languages") + ": ")
        run.bold = True
        lang_strs = []
        for entry in languages.get("entries", []):
            lang_strs.append(f"{entry['language']} ({entry.get('fluency', '')})")
        p.add_run(" | ".join(lang_strs))


def add_work_section(doc: Document, data: dict, styled: bool) -> None:
    """Render work experience entries."""
    work = data.get("work")
    if not work:
        return
    lang = data.get("meta", {}).get("language", "en")
    doc.add_heading(work.get("title", "Work Experience"), level=1)

    for job in work.get("entries", []):
        company = job.get("company", "")
        position = job.get("position", "")
        start = job.get("start_date", "")
        end = job.get("end_date", "")
        date_range = fmt_date_range(start, end, lang)
        technologies = job.get("technologies", [])

        # Logo (styled only)
        container = doc
        if styled:
            logo_path = resolve_logo(job.get("logo", ""))
            
            table = doc.add_table(rows=1, cols=2)
            remove_table_borders(table)
            
            logo_cell = table.cell(0, 0)
            set_cell_width(logo_cell, 800)  # ~1.25 cm
            
            if logo_path:
                lp = logo_cell.paragraphs[0]
                lp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = lp.add_run()
                try:
                    # make it slightly smaller to fit better in left margin context
                    run.add_picture(str(logo_path), width=Cm(0.9))
                except Exception:
                    pass  # Skip if image format is unsupported
            
            container = table.cell(0, 1)
            p = container.paragraphs[0]
        else:
            p = doc.add_paragraph()

        # Position + date
        run = p.add_run(position)
        run.bold = True
        run.font.size = Pt(11)
        if date_range:
            p.add_run(f"  |  {date_range}").font.color.rgb = LIGHT_GRAY

        # Company
        p = container.add_paragraph()
        run = p.add_run(company)
        run.font.color.rgb = ACCENT if styled else DARK
        run.bold = True
        run.font.size = Pt(10)

        # Technologies
        if technologies:
            p = container.add_paragraph()
            run = p.add_run(", ".join(technologies))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY

        # Summary
        summary = job.get("summary", "")
        if summary:
            container.add_paragraph(clean_text(summary))

        # Highlights
        for h in job.get("highlights", []):
            container.add_paragraph(clean_text(h), style="List Bullet")

        # Recommendations
        for rec in job.get("recommendations", []):
            quote = rec.get("quote", "")
            if quote:
                p = container.add_paragraph(style="Quote")
                p.add_run(f"\u201C{clean_text(quote)}\u201D")
                p = container.add_paragraph()
                author = rec.get("author", "")
                role = rec.get("role", "")
                run = p.add_run(f"\u2014 {author}, {role}")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = DARK
                
        if styled:
            doc.add_paragraph()


def add_products_section(doc: Document, data: dict, styled: bool) -> None:
    """Render own products & innovations."""
    products = data.get("products")
    if not products:
        return
    lang = data.get("meta", {}).get("language", "en")
    doc.add_heading(products.get("title", "Products"), level=1)

    for prod in products.get("entries", []):
        name = prod.get("name", "")
        start = prod.get("start_date", "")
        end = prod.get("end_date", "")
        date_range = fmt_year_range(start, end, lang)
        technologies = prod.get("technologies", [])

        # Logo (styled only)
        container = doc
        if styled:
            logo_path = resolve_logo(prod.get("logo", ""))
            
            table = doc.add_table(rows=1, cols=2)
            remove_table_borders(table)
            
            logo_cell = table.cell(0, 0)
            set_cell_width(logo_cell, 800)  # ~1.25 cm
            
            if logo_path:
                lp = logo_cell.paragraphs[0]
                lp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = lp.add_run()
                try:
                    run.add_picture(str(logo_path), width=Cm(0.9))
                except Exception:
                    pass
            
            container = table.cell(0, 1)
            p = container.paragraphs[0]
        else:
            p = doc.add_paragraph()

        # Name + date
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(11)
        if date_range:
            p.add_run(f"  |  {date_range}").font.color.rgb = LIGHT_GRAY

        # Technologies
        if technologies:
            p = container.add_paragraph()
            run = p.add_run(", ".join(technologies))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY

        # URL (styled only)
        if styled:
            url = prod.get("url", "")
            if url:
                p = container.add_paragraph()
                run = p.add_run(url)
                run.font.size = Pt(9)
                run.font.color.rgb = ACCENT

        # Summary
        summary = prod.get("summary", "")
        if summary:
            container.add_paragraph(clean_text(summary))

        # Highlights
        for h in prod.get("highlights", []):
            container.add_paragraph(clean_text(h), style="List Bullet")
            
        if styled:
            doc.add_paragraph()


def add_education_section(doc: Document, data: dict, styled: bool) -> None:
    """Render education entries."""
    education = data.get("education")
    if not education:
        return
    lang = data.get("meta", {}).get("language", "en")
    doc.add_heading(education.get("title", "Education"), level=1)

    for edu in education.get("entries", []):
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        start = edu.get("start_date", "")
        end = edu.get("end_date", "")
        date_range = fmt_year_range(start, end, lang)

        p = doc.add_paragraph()
        run = p.add_run(degree)
        run.bold = True
        if date_range:
            p.add_run(f"  |  {date_range}").font.color.rgb = LIGHT_GRAY

        p = doc.add_paragraph()
        run = p.add_run(institution)
        run.font.color.rgb = ACCENT if styled else DARK
        field = edu.get("field", "")
        if field:
            p.add_run(f" | {field}")

        desc = edu.get("description", "")
        if desc:
            doc.add_paragraph(clean_text(desc))


def add_publications_section(doc: Document, data: dict, styled: bool) -> None:
    """Render publications."""
    publications = data.get("publications")
    if not publications:
        return
    doc.add_heading(publications.get("title", "Publications"), level=1)

    for pub in publications.get("entries", []):
        p = doc.add_paragraph()
        role = pub.get("role", "")
        title = pub.get("title", "")
        run = p.add_run(f"{role}: ")
        run.bold = True
        run = p.add_run(title)
        run.italic = True
        publisher = pub.get("publisher", "")
        date = pub.get("date", "")
        if publisher or date:
            p.add_run(f" | {publisher}, {date}")
        desc = pub.get("description", "")
        if desc:
            doc.add_paragraph(clean_text(desc))


def add_training_section(doc: Document, data: dict, styled: bool) -> None:
    """Render training & speaking engagements."""
    training = data.get("training")
    if not training:
        return
    lang = data.get("meta", {}).get("language", "en")
    doc.add_heading(training.get("title", "Training & Speaking"), level=1)

    for t in training.get("entries", []):
        title = t.get("title", "")
        org = t.get("organization", "")
        location = t.get("location", "")

        if t.get("start_date"):
            date_str = fmt_date_range(t["start_date"], t.get("end_date", ""), lang)
        elif t.get("date"):
            date_str = fmt_date(t["date"], lang)
        else:
            date_str = ""

        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        if date_str:
            p.add_run(f"  |  {date_str}").font.color.rgb = LIGHT_GRAY

        p = doc.add_paragraph()
        run = p.add_run(org)
        run.font.color.rgb = ACCENT if styled else DARK
        if location:
            p.add_run(f" | {location}")

        desc = t.get("description", "")
        if desc:
            doc.add_paragraph(clean_text(desc))


def add_about_section(doc: Document, data: dict, styled: bool) -> None:
    """Render about me / interests section."""
    about = data.get("about")
    if not about:
        return
    doc.add_heading(about.get("title", "About Me"), level=1)
    text = about.get("text", "")
    if text:
        doc.add_paragraph(clean_text(text))

    interests = about.get("interests", [])
    if interests:
        for interest in interests:
            p = doc.add_paragraph()
            run = p.add_run(interest.get("title", ""))
            run.bold = True
            desc = interest.get("description", "")
            if desc:
                p.add_run(f" \u2014 {clean_text(desc)}")


def add_open_source_section(doc: Document, data: dict, styled: bool) -> None:
    """Render open source contributions."""
    oss = data.get("open_source")
    if not oss:
        return
    doc.add_heading(oss.get("title", "Open Source"), level=1)
    for proj in oss.get("entries", []):
        p = doc.add_paragraph()
        run = p.add_run(proj.get("name", ""))
        run.bold = True
        desc = proj.get("description", "")
        if desc:
            p.add_run(f" \u2014 {clean_text(desc)}")


def add_cta_section(doc: Document, data: dict, styled: bool) -> None:
    """Render call-to-action footer."""
    cta = data.get("cta")
    if not cta:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    text = cta.get("text", "")
    run = p.add_run(clean_text(text))
    run.italic = True
    run.font.color.rgb = ACCENT if styled else DARK

    contact = f"{cta.get('email', '')} {cta.get('or', 'or')} {cta.get('phone', '')}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(contact)
    run.font.color.rgb = ACCENT if styled else DARK
    run.bold = True


def add_footer(doc: Document, data: dict) -> None:
    """Add last-updated footer."""
    meta = data.get("meta", {})
    updated = meta.get("last_updated", "")
    url = meta.get("footer_url_text", "")
    if updated or url:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        parts = []
        if url:
            parts.append(url)
        if updated:
            label = meta.get("footer_date_label", "Last updated")
            parts.append(f"{label}: {updated}")
        run = p.add_run(" | ".join(parts))
        run.font.size = Pt(8)
        run.font.color.rgb = LIGHT_GRAY


# ──────────────────────────────────────────────
# SECTION RENDERERS — SUMMARY
# ──────────────────────────────────────────────
def add_header_summary(doc: Document, data: dict, styled: bool) -> None:
    """Render header for summary YAML: photo left, name/contact right, accent rule."""
    basics = data["basics"]

    address = basics.get("address", "")
    if isinstance(address, dict):
        address = address.get("full", "")
    company = basics.get("company", "")
    title = basics.get("title", "")
    label = f"{company}\u00AE | {title}" if company else title

    contact_parts = [
        address,
        basics.get("email", ""),
        basics.get("phone", ""),
        basics.get("website", {}).get("text", ""),
    ]
    profiles = basics.get("profiles", [])
    profile_texts = []
    for pr in profiles:
        network = pr.get("network", "")
        text = pr.get("text", pr.get("url", ""))
        profile_texts.append(f"{network}: {text}" if network else text)

    if styled:
        photo_path = resolve_asset(basics.get("photo", ""))

        # ── Two-column table: [photo | name + contact] ──
        table = doc.add_table(rows=1, cols=2 if photo_path else 1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        remove_table_borders(table)

        if photo_path:
            # Left cell: photo
            photo_cell = table.cell(0, 0)
            set_cell_vertical_alignment(photo_cell, "center")
            set_cell_width(photo_cell, 1600)  # ~2.8 cm
            pp = photo_cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = pp.add_run()
            run.add_picture(str(photo_path), width=Cm(2.5))

            text_cell = table.cell(0, 1)
        else:
            text_cell = table.cell(0, 0)

        set_cell_vertical_alignment(text_cell, "center")

        # Name
        np = text_cell.paragraphs[0]
        np.paragraph_format.space_after = Pt(1)
        run = np.add_run(basics["name"])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DARK

        # Company | Title
        tp = text_cell.add_paragraph()
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(3)
        run = tp.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT

        # Contact info line
        contact_line = " | ".join(part for part in contact_parts if part)
        cp = text_cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(1)
        run = cp.add_run(contact_line)
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

        # Profiles
        if profile_texts:
            prp = text_cell.add_paragraph()
            prp.paragraph_format.space_before = Pt(0)
            prp.paragraph_format.space_after = Pt(0)
            run = prp.add_run(" | ".join(profile_texts))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY

        # Accent rule below header
        add_accent_rule(doc)

    else:
        # ── ATS: plain centered text ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(basics["name"])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = DARK

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(11)
        run.font.color.rgb = MID_GRAY

        contact_line = " | ".join(part for part in contact_parts if part)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact_line)
        run.font.size = Pt(9)
        run.font.color.rgb = MID_GRAY

        if profile_texts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" | ".join(profile_texts))
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY


def add_summary_section_brief(doc: Document, data: dict, styled: bool) -> None:
    """Render summary text from the summary YAML (separate 'summary' key)."""
    summary = data.get("summary", {})
    if not summary:
        return
    doc.add_heading(summary.get("title", "Summary"), level=1)
    text = summary.get("text", "")
    if text:
        doc.add_paragraph(clean_text(text))


def add_tech_stack_section(doc: Document, data: dict, styled: bool) -> None:
    """Render tech stack (summary format: label + content string)."""
    tech = data.get("tech_stack")
    if not tech:
        return
    doc.add_heading(tech.get("title", "Technical Stack"), level=1)
    for cat in tech.get("categories", []):
        p = doc.add_paragraph()
        run = p.add_run(cat.get("label", ""))
        run.bold = True
        p.add_run(f": {cat.get('content', '')}")


def add_clients_section(doc: Document, data: dict, styled: bool) -> None:
    """Render notable clients (summary format)."""
    clients = data.get("clients")
    if not clients:
        return
    doc.add_heading(clients.get("title", "Clients"), level=1)
    for entry in clients.get("entries", []):
        title = entry.get("title", "")
        years = entry.get("years", "")
        p = doc.add_paragraph()
        run = p.add_run(clean_text(title))
        run.bold = True
        if years:
            years_clean = years.replace("--", "\u2013")
            p.add_run(f"  |  {years_clean}").font.color.rgb = LIGHT_GRAY
        desc = entry.get("description", "")
        if desc:
            doc.add_paragraph(clean_text(desc))


def add_products_section_summary(doc: Document, data: dict, styled: bool) -> None:
    """Render products in summary format (title + years + description)."""
    products = data.get("products")
    if not products:
        return
    doc.add_heading(products.get("title", "Products"), level=1)
    for entry in products.get("entries", []):
        title = entry.get("title", "")
        years = entry.get("years", "")
        p = doc.add_paragraph()
        run = p.add_run(clean_text(title))
        run.bold = True
        if years:
            years_clean = years.replace("--", "\u2013")
            p.add_run(f"  |  {years_clean}").font.color.rgb = LIGHT_GRAY
        desc = entry.get("description", "")
        if desc:
            doc.add_paragraph(clean_text(desc))


def add_education_section_summary(doc: Document, data: dict, styled: bool) -> None:
    """Render education in summary format (flat structure)."""
    education = data.get("education")
    if not education:
        return
    doc.add_heading(education.get("title", "Education"), level=1)

    degree = education.get("degree", "")
    institution = education.get("institution", "")
    years = education.get("years", "").replace("--", "\u2013")
    p = doc.add_paragraph()
    run = p.add_run(degree)
    run.bold = True
    if years:
        p.add_run(f"  |  {years}").font.color.rgb = LIGHT_GRAY
    if institution:
        p = doc.add_paragraph()
        run = p.add_run(institution)
        run.font.color.rgb = ACCENT if styled else DARK

    # Publication
    pub = education.get("publication")
    if pub:
        p = doc.add_paragraph()
        run = p.add_run(f"{pub.get('role', '')}: ")
        run.bold = True
        run = p.add_run(pub.get("title", ""))
        run.italic = True
        publisher = pub.get("publisher", "")
        year = pub.get("year", "")
        if publisher or year:
            p.add_run(f" ({publisher}, {year})")

    speaking = education.get("speaking", "")
    if speaking:
        doc.add_paragraph(clean_text(speaking))


# ──────────────────────────────────────────────
# DOCUMENT BUILDERS
# ──────────────────────────────────────────────
def build_resume_docx(data: dict, styled: bool) -> Document:
    """Build a DOCX document from full resume YAML data."""
    doc = Document()

    # Set page margins — smaller top margin for styled so banner sits near top
    for section in doc.sections:
        section.top_margin = Cm(0.8) if styled else Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    setup_styles(doc, styled)
    add_header_resume(doc, data, styled)
    add_summary_section(doc, data, styled)
    add_services_section(doc, data, styled)
    add_skills_section(doc, data, styled)
    add_work_section(doc, data, styled)
    add_products_section(doc, data, styled)
    add_education_section(doc, data, styled)
    add_publications_section(doc, data, styled)
    add_training_section(doc, data, styled)
    add_about_section(doc, data, styled)
    add_open_source_section(doc, data, styled)
    add_cta_section(doc, data, styled)
    add_footer(doc, data)
    return doc


def build_summary_docx(data: dict, styled: bool) -> Document:
    """Build a DOCX document from summary YAML data."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    setup_styles(doc, styled)
    add_header_summary(doc, data, styled)
    add_summary_section_brief(doc, data, styled)
    add_services_section(doc, data, styled)
    add_tech_stack_section(doc, data, styled)
    add_clients_section(doc, data, styled)
    add_products_section_summary(doc, data, styled)
    add_education_section_summary(doc, data, styled)
    add_cta_section(doc, data, styled)
    add_footer(doc, data)
    return doc


# ──────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────
def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("data_file", help="Path to resume YAML data file")
    parser.add_argument("output_file", help="Path for DOCX output")
    parser.add_argument(
        "--style",
        choices=["ats", "styled"],
        default="styled",
        help="Visual style: 'ats' (clean, no images) or 'styled' (accent colors, photos, logos)",
    )
    args = parser.parse_args()

    data_path = Path(args.data_file)
    out_path = Path(args.output_file)

    if not data_path.exists():
        print(f"Error: {data_path} not found", file=sys.stderr)
        sys.exit(1)

    with data_path.open("r", encoding="utf-8") as fh:
        data = yaml.safe_load(fh)

    styled = args.style == "styled"

    if is_resume_data(data):
        doc = build_resume_docx(data, styled)
    else:
        doc = build_summary_docx(data, styled)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
